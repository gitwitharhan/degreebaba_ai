from __future__ import annotations

from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import BinaryIO

from docx import Document


def read_docx(file: str | Path | BinaryIO) -> dict:
    """Read a DOCX and return paragraphs plus tables without summarising content."""
    document = Document(file)

    paragraphs = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        paragraphs.append(
            {
                "text": text,
                "style": paragraph.style.name if paragraph.style else "",
            }
        )

    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append(rows)

    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "raw_text": "\n".join(item["text"] for item in paragraphs),
    }


async def read_upload_file(upload_file) -> dict:
    """FastAPI UploadFile adapter for python-docx."""
    suffix = Path(upload_file.filename or "").suffix or ".docx"
    with NamedTemporaryFile(suffix=suffix) as temp:
        temp.write(await upload_file.read())
        temp.flush()
        return read_docx(temp.name)
